
from datetime import datetime
import streamlit as st
import json
import google.generativeai as genai
from docx import Document
from io import BytesIO

# ==========================================
# CONFIGURATION
# ==========================================



GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("models/gemini-2.5-flash")

# ==========================================
# PAGE SETTINGS
# ==========================================

st.set_page_config(
    page_title="AI SDLC Copilot",
    page_icon="🏦",
    layout="wide"
)

st.title("🏦 AI SDLC Copilot")
st.markdown("Accelerate the Software Development Lifecycle with AI-powered Requirements, Testing, SQL, and Documentation.")

# ==========================================
# WORD EXPORT FUNCTIONS
# ==========================================

def create_word_file(content, title):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_professional_brd(data):

    doc = Document()

    doc.add_heading("Business Requirements Document", level=0)
    doc.add_paragraph("AI SDLC Copilot")
    doc.add_paragraph(
        f"Generated On: {datetime.today().strftime('%d-%b-%Y')}"
    )

    doc.add_page_break()

    doc.add_heading("Business Objective", level=1)
    for item in data.get("business_objective", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("Project Scope", level=1)
    for item in data.get("project_scope", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("Stakeholders", level=1)
    stakeholder_table = doc.add_table(rows=1, cols=2)
    stakeholder_table.style = "Table Grid"
    stakeholder_table.cell(0, 0).text = "Stakeholder"
    stakeholder_table.cell(0, 1).text = "Role"

    for stakeholder in data.get("stakeholders", []):
        row = stakeholder_table.add_row().cells
        row[0].text = str(stakeholder)
        row[1].text = "Business User"

    doc.add_heading("Business Requirements", level=1)
    br_table = doc.add_table(rows=1, cols=2)
    br_table.style = "Table Grid"
    br_table.cell(0, 0).text = "BR ID"
    br_table.cell(0, 1).text = "Requirement"

    for i, req in enumerate(data.get("business_requirements", []), start=1):
        row = br_table.add_row().cells
        row[0].text = f"BR-{i:03d}"
        row[1].text = str(req)

    doc.add_heading("Functional Requirements", level=1)
    fr_table = doc.add_table(rows=1, cols=2)
    fr_table.style = "Table Grid"
    fr_table.cell(0, 0).text = "FR ID"
    fr_table.cell(0, 1).text = "Requirement"

    for i, req in enumerate(data.get("functional_requirements", []), start=1):
        row = fr_table.add_row().cells
        row[0].text = f"FR-{i:03d}"
        row[1].text = str(req)

    doc.add_heading("Assumptions", level=1)
    for item in data.get("assumptions", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("Dependencies", level=1)
    for item in data.get("dependencies", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("Risks", level=1)
    for item in data.get("risks", []):
        doc.add_paragraph(str(item), style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from datetime import datetime


def create_professional_user_story(data):

    doc = Document()

    # =====================================================
    # COVER PAGE
    # =====================================================

    title = doc.add_heading("USER STORY DOCUMENT", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("AI SDLC Copilot\n").bold = True
    p.add_run(
        f"Generated On: {datetime.today().strftime('%d-%b-%Y')}"
    )

    doc.add_page_break()

    # =====================================================
    # DOCUMENT INFORMATION
    # =====================================================

    doc.add_heading("Document Information", level=1)

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"

    info.cell(0,0).text = "Project"
    info.cell(0,1).text = data.get("project_name","")

    info.cell(1,0).text = "Epic"
    info.cell(1,1).text = data.get("epic","")

    info.cell(2,0).text = "Version"
    info.cell(2,1).text = "1.0"

    info.cell(3,0).text = "Generated Date"
    info.cell(3,1).text = datetime.today().strftime("%d-%b-%Y")

    # =====================================================
    # USER STORIES
    # =====================================================

    for story in data.get("stories", []):

        doc.add_page_break()

        doc.add_heading(story.get("story_id",""), level=1)

        # ------------------------
        # Story Summary
        # ------------------------

        doc.add_heading("Story Summary", level=2)

        summary = doc.add_table(rows=5, cols=2)
        summary.style = "Table Grid"

        summary.cell(0,0).text = "Feature"
        summary.cell(0,1).text = story.get("feature","")

        summary.cell(1,0).text = "Priority"
        summary.cell(1,1).text = story.get("priority","")

        summary.cell(2,0).text = "Story Points"
        summary.cell(2,1).text = str(story.get("story_points",""))

        summary.cell(3,0).text = "Sprint"
        summary.cell(3,1).text = story.get("sprint","")

        summary.cell(4,0).text = "Story ID"
        summary.cell(4,1).text = story.get("story_id","")

        # ------------------------
        # User Story
        # ------------------------

        doc.add_heading("User Story", level=2)
        doc.add_paragraph(story.get("user_story",""))

        # ------------------------
        # Business Value
        # ------------------------

        doc.add_heading("Business Value", level=2)
        doc.add_paragraph(story.get("business_value",""))

        # ------------------------
        # Acceptance Criteria
        # ------------------------

        doc.add_heading("Acceptance Criteria", level=2)

        for ac in story.get("acceptance_criteria", []):
            doc.add_paragraph(ac, style="List Bullet")

        # ------------------------
        # Definition of Done
        # ------------------------

        doc.add_heading("Definition of Done", level=2)

        for dod in story.get("definition_of_done", []):
            doc.add_paragraph(dod, style="List Bullet")

        # ------------------------
        # Dependencies
        # ------------------------

        doc.add_heading("Dependencies", level=2)

        for dep in story.get("dependencies", []):
            doc.add_paragraph(dep, style="List Bullet")

        # ------------------------
        # Assumptions
        # ------------------------

        doc.add_heading("Assumptions", level=2)

        for ass in story.get("assumptions", []):
            doc.add_paragraph(ass, style="List Bullet")

    # =====================================================
    # APPROVAL
    # =====================================================

    doc.add_page_break()

    doc.add_heading("Approval", level=1)

    approval = doc.add_table(rows=2, cols=3)
    approval.style = "Table Grid"

    approval.cell(0,0).text = "Role"
    approval.cell(0,1).text = "Name"
    approval.cell(0,2).text = "Signature"

    approval.cell(1,0).text = "Business Owner"

    # =====================================================

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


from docx import Document
from io import BytesIO
from datetime import datetime
import re

def create_professional_uat(content):

    doc = Document()

    # =========================
    # COVER PAGE
    # =========================
    doc.add_heading("UAT TEST DOCUMENT", 0)
    doc.add_paragraph("AI SDLC Copilot")
    doc.add_paragraph(f"Generated On: {datetime.today().strftime('%d-%b-%Y')}")
    doc.add_page_break()

    # =========================
    # RAW SUMMARY SECTION
    # =========================
    doc.add_heading("1. Test Summary", level=1)
    doc.add_paragraph("Generated UAT scenarios based on business requirements.")

    # =========================
    # TEST CASES SECTION
    # =========================
    doc.add_heading("2. UAT Test Cases", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Test Case ID"
    hdr[1].text = "Scenario"
    hdr[2].text = "Expected Result"
    hdr[3].text = "Type"

    # =========================
    # SMART PARSING (basic AI output handling)
    # =========================

    lines = content.split("\n")
    tc_id = 1

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Try to detect test case style lines
        if any(keyword in line.lower() for keyword in ["test", "verify", "ensure", "check", "validate"]):

            row = table.add_row().cells
            row[0].text = f"TC-{tc_id:03d}"
            row[1].text = line[:120]

            row[2].text = "As per requirement"
            row[3].text = "Positive/Negative"

            tc_id += 1

    # fallback if no parsing worked
    if tc_id == 1:
        row = table.add_row().cells
        row[0].text = "TC-001"
        row[1].text = content[:200]
        row[2].text = "To be validated as per requirement"
        row[3].text = "General"

    # =========================
    # SIGN OFF SECTION
    # =========================
    doc.add_heading("3. Sign Off", level=1)
    doc.add_paragraph("Business Analyst: __________________")
    doc.add_paragraph("QA Lead: __________________")
    doc.add_paragraph("Business Owner: __________________")

    # =========================
    # EXPORT
    # =========================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
    
    
import pandas as pd
from io import BytesIO

def create_professional_uat_excel(data):

    rows = []

    for tc in data.get("test_cases", []):

        rows.append({
            "Test Case ID": tc.get("id", ""),
            "Module": tc.get("module", ""),
            "Scenario": tc.get("scenario", ""),
            "Priority": tc.get("priority", ""),
            "Preconditions": tc.get("preconditions", ""),
            "Test Steps": tc.get("test_steps", ""),
            "Expected Result": tc.get("expected_result", "")
        })

    df = pd.DataFrame(rows)

    buffer = BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="UAT Test Cases")

    buffer.seek(0)
    return buffer


def create_professional_sql(content):

    doc = Document()

    doc.add_heading("SQL Query Document", level=0)

    doc.add_paragraph("AI SDLC Copilot")

    doc.add_paragraph(
        f"Generated On: {datetime.today().strftime('%d-%b-%Y')}"
    )

    doc.add_page_break()

    doc.add_heading("Generated SQL Query", level=1)

    doc.add_paragraph(content)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer
    
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
def create_professional_technical_test(data):

    doc = Document()

    # -------------------------
    # Cover Page
    # -------------------------

    title = doc.add_heading("Technical Test Case Document", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("AI SDLC Copilot\n").bold = True
    p.add_run(f"Generated On: {datetime.today().strftime('%d-%b-%Y')}")

    doc.add_page_break()

    # -------------------------
    # Document Information
    # -------------------------

    doc.add_heading("Document Information", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    table.cell(0,0).text = "Application"
    table.cell(0,1).text = data["application"]

    table.cell(1,0).text = "Module"
    table.cell(1,1).text = data["module"]

    table.cell(2,0).text = "Version"
    table.cell(2,1).text = data["version"]

    table.cell(3,0).text = "Generated By"
    table.cell(3,1).text = "AI SDLC Copilot"

    doc.add_page_break()
    
    for tc in data["test_cases"]:

        doc.add_heading(tc["test_case_id"], level=1)

        summary = doc.add_table(rows=5, cols=2)
        summary.style = "Table Grid"

        summary.cell(0,0).text = "Component"
        summary.cell(0,1).text = tc["component"]

        summary.cell(1,0).text = "Scenario"
        summary.cell(1,1).text = tc["scenario"]

        summary.cell(2,0).text = "Priority"
        summary.cell(2,1).text = tc["priority"]

        summary.cell(3,0).text = "Test Type"
        summary.cell(3,1).text = tc["test_type"]

        summary.cell(4,0).text = "Automation"
        summary.cell(4,1).text = tc["automation"]
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

# ==========================================
# TABS
# ==========================================
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
    [
        "📄 Requirements Document Generator",
        "📖 User Story Generator",
        "🧪 UAT Test Case Generator",
        "💻 SQL Query Generator",
        "🧩 Technical Test Case Generator",
        "📝 Prompt Templates"
    ]
)
# ==========================================
# Requirements Document Generator
# ==========================================

with tab1:

    brd_input = st.text_area("Paste Requirement", height=250)

    if st.button("Generate BRD"):

        prompt = f"""
Return ONLY valid JSON.

{{
"business_objective":[],
"project_scope":[],
"stakeholders":[],
"business_requirements":[],
"functional_requirements":[],
"assumptions":[],
"dependencies":[],
"risks":[]
}}

Requirement:
{brd_input}
"""

        response = model.generate_content(prompt)

        json_text = response.text
        json_text = json_text.replace("```json", "")
        json_text = json_text.replace("```", "")

        json_text = response.text

        st.code(json_text)  # DEBUG (IMPORTANT)

        json_text = json_text.replace("```json", "").replace("```", "").strip()

        data = json.loads(json_text)

        st.json(data)

        doc_file = create_professional_brd(data)

        st.download_button(
            "📥 Download BRD",
            data=doc_file,
            file_name="Credit_Risk_BRD.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ==========================================
# USER STORY
# ==========================================

with tab2:

    story_input = st.text_area("Requirement", key="story")

    if st.button("Generate User Stories"):

        prompt = f"""
        You are a Senior Agile Business Analyst.

        Return ONLY valid JSON.

        {{
          "project_name": "",
          "epic": "",
          "stories": [
            {{
              "story_id": "US-001",
              "feature": "",
              "priority": "",
              "story_points": "",
              "sprint": "",
              "user_story": "",
              "business_value": "",
              "acceptance_criteria": [],
              "definition_of_done": [],
              "dependencies": [],
              "assumptions": []
            }}
          ]
        }}

        Requirement:

        {story_input}
        """

        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.2,
                "response_mime_type": "application/json"
            }
)

        json_text = response.text
        json_text = json_text.replace("```json","").replace("```","").strip()

        data = json.loads(json_text)

        doc_file = create_professional_user_story(data)

        st.download_button(
            "📥 Download User Stories",
            data=doc_file,
            file_name="User_Stories.docx"
        )

# ==========================================
# UAT
# ==========================================

with tab3:


    uat_input = st.text_area("Requirement", key="uat")

    if st.button("Generate UAT"):

        response = model.generate_content(
            f"Generate detailed UAT test cases for:\n\n{uat_input}"
        )

        output = response.text

        st.markdown(output)

        doc_file = create_professional_uat(output)

        st.download_button(
            "📥 Download UAT",
            data=doc_file,
            file_name="UAT_Test_Cases.docx"
        )

# ==========================================
# SQL
# ==========================================

with tab4:

    sql_input = st.text_area("Requirement", key="sql")

    if st.button("Generate SQL"):

        response = model.generate_content(
            f"Generate SQL query, explanation and assumptions for: {sql_input}"
        )

        output = response.text
        st.markdown(output)

        doc_file = create_professional_sql(output)

        st.download_button(
            "📥 Download SQL",
            data=doc_file,
            file_name="SQL_Output.docx"
        )
        
        
            # ==========================================
        # PROMPT TEMPLATES
        # ==========================================

# ==========================================
# PROMPT TEMPLATES
# ==========================================

with tab6:

    st.subheader("📝 BA Prompt Templates Library")

    prompt_type = st.selectbox(
    "Select Template",
    [
        "BRD Prompt",
        "User Story Prompt",
        "UAT Prompt",
        "Technical Test Case Prompt",
        "SQL Prompt"
    ]
)

    prompt_templates = {

        "BRD Prompt": """
You are a Senior Banking Business Analyst.

Generate a detailed Business Requirements Document.

Return ONLY valid JSON.

{
  "business_objective": [],
  "project_scope": [],
  "stakeholders": [],
  "business_requirements": [],
  "functional_requirements": [],
  "assumptions": [],
  "dependencies": [],
  "risks": []
}

Context:
Credit Risk, Watchlist, UTP, Forbearance, Impairment,
Regulatory Reporting and Banking.

Stakeholder Notes:

{stakeholder_notes}
""",

        "User Story Prompt": """
You are a Senior Agile Business Analyst.

Generate:

1. Epic
2. Features
3. User Stories

Format:

As a <role>
I want <goal>
So that <benefit>

Also provide:

- Acceptance Criteria
- Definition of Done

Requirement:

{requirement}
""",

        "UAT Prompt": """
You are a QA Lead.

Generate detailed UAT Test Cases.

Include:

- Test Case ID
- Test Scenario
- Preconditions
- Test Steps
- Expected Results

Cover:

- Positive Scenarios
- Negative Scenarios
- Boundary Scenarios

Requirement:

{requirement}
""",

        "SQL Prompt": """
You are an SQL expert.

Generate:

1. SQL Query
2. Query Explanation
3. Assumptions
4. Optimization Suggestions

Requirement:

{requirement}
""",
        "Technical Test Case Prompt": """
You are a Senior Software Test Architect.

Generate detailed Technical Test Cases.

Return ONLY valid JSON.

{
  "application": "",
  "module": "",
  "version": "1.0",
  "test_cases": [
    {
      "test_case_id": "",
      "component": "",
      "scenario": "",
      "priority": "",
      "test_type": "",
      "preconditions": [],
      "test_steps": [],
      "expected_result": "",
      "negative_tests": [],
      "boundary_tests": [],
      "exception_handling": [],
      "automation": ""
    }
  ]
}

Requirement:

{requirement}
"""
    }

    output_templates = {

        "BRD Prompt": """
{
  "business_objective": [],
  "project_scope": [],
  "stakeholders": [],
  "business_requirements": [],
  "functional_requirements": [],
  "assumptions": [],
  "dependencies": [],
  "risks": []
}
""",

        "User Story Prompt": """
Epic

Features

User Stories

Acceptance Criteria

Definition of Done
""",

        "UAT Prompt": """
TC-001

Scenario

Preconditions

Steps

Expected Result
""",

        "SQL Prompt": """
SELECT customer_id,
       customer_name,
       status
FROM customer
WHERE status = 'ACTIVE';
""",
"Technical Test Case Prompt": """
{
  "application": "Credit Risk System",
  "module": "Watchlist",
  "version": "1.0",
  "test_cases": [
    {
      "test_case_id": "TC-001",
      "component": "Watchlist Service",
      "scenario": "Add customer to Watchlist",
      "priority": "High",
      "test_type": "Functional",
      "preconditions": [
        "User is logged in",
        "Customer exists"
      ],
      "test_steps": [
        "Navigate to Watchlist",
        "Click Add",
        "Enter Customer ID",
        "Click Save"
      ],
      "expected_result": "Customer is added successfully.",
      "negative_tests": [
        "Duplicate customer",
        "Invalid customer ID"
      ],
      "boundary_tests": [
        "Maximum remarks length",
        "Empty remarks"
      ],
      "exception_handling": [
        "Database unavailable",
        "API timeout"
      ],
      "automation": "Yes"
    }
  ]
}
"""
    }

    selected_prompt = prompt_templates[prompt_type]
    selected_output = output_templates[prompt_type]

    st.markdown("## 📌 Prompt Used by AI")

    st.code(
        selected_prompt,
        language="text"
    )

    st.markdown("## 📌 Expected Output Structure")

    st.code(
        selected_output,
        language="text"
    )

    st.download_button(
        label="📥 Download Prompt Template",
        data=selected_prompt,
        file_name=f"{prompt_type.replace(' ', '_')}.txt",
        mime="text/plain"
    )
# ==========================================
# Technical Test Case Generator
# ==========================================

with tab5:

    st.subheader("🧩 Technical Test Case Generator")

    tech_input = st.text_area(
        "Enter Requirement",
        key="technical_test"
    )

    if st.button("Generate Technical Test Cases"):

        prompt = f"""
You are a Senior Software Test Architect.

Return ONLY valid JSON.

{{
    "application":"",
    "module":"",
    "version":"1.0",
    "test_cases":[
        {{
            "test_case_id":"",
            "component":"",
            "scenario":"",
            "priority":"",
            "test_type":"",
            "preconditions":[],
            "test_steps":[],
            "expected_result":"",
            "negative_tests":[],
            "boundary_tests":[],
            "exception_handling":[],
            "automation":""
        }}
    ]
}}

Requirement:

{tech_input}
"""

        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.2,
                "response_mime_type": "application/json"
            }
        )

        json_text = response.text.strip()

        try:
            data = json.loads(json_text)

            st.success("Technical Test Cases Generated")
            st.json(data)
            
            doc_file = create_professional_technical_test(data)

            st.download_button(
                "📥 Download Technical Test Cases",
                data=doc_file,
                file_name="Technical_Test_Cases.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error("Invalid JSON returned by Gemini")
            st.code(json_text)
            st.write(e)